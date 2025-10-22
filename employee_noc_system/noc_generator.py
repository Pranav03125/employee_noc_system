import os
import re
from docx import Document

class EmployeeNOCGenerator:
    """
    Generates personalized NOC documents by inserting user input next to the placeholders:
      - "Full Name:"
      - "Job Title:"
      - "Department:"
    The replacement keeps the placeholder label visible and places the user input immediately after the colon,
    e.g. "Full Name: Arjun Kumar".
    """

    def __init__(self, template_path: str):
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
        self.template_path = template_path

    def _replace_fields_in_text(self, text: str, replacements: dict) -> str:
        """
        Replace any of the keys in `replacements` when they appear as 'Key: ...' (case-insensitive).
        Match only up to the next newline so replacing one label won't remove subsequent labels that are on
        separate lines inside the same paragraph.
        """
        updated = text
        for key, value in replacements.items():
            pattern = re.compile(rf'({re.escape(key)})\s*:\s*([^\n]*)', re.IGNORECASE)
            updated = pattern.sub(lambda m: f"{m.group(1)}: {value}", updated)
        return updated

    def _find_label_match(self, text: str, key: str):
        """Return a regex match for the label (e.g. 'Job Title:') in text, or None."""
        pattern = re.compile(rf'\b{re.escape(key)}\b\s*:', re.IGNORECASE)
        return pattern.search(text)

    def _replace_field_in_paragraph_runs(self, para, key: str, value: str) -> bool:
        """
        Replace only the value for a given label inside a paragraph while preserving run objects.
        This replaces text after the label's colon up to the next newline (so later placeholders on
        subsequent lines are preserved).
        Returns True if a replacement was made.
        """
        full_text = para.text
        match = self._find_label_match(full_text, key)
        if not match:
            return False

        label_end = match.end()  # index right after the colon
        # find end of current line (next newline) to avoid touching following placeholders
        next_nl = full_text.find('\n', label_end)
        if next_nl == -1:
            rest = full_text[label_end:]  # no newline, keep whatever follows
        else:
            rest = full_text[next_nl:]   # include newline and everything after it

        # Build the desired paragraph text: original upto label_end + space + value + rest
        left_part = full_text[:label_end]
        desired_text = f"{left_part} {value}{rest}"

        # Write desired_text back into runs, preserving run objects (and inline drawings)
        cum_len = 0
        runs = para.runs
        # If there are no runs, fallback to para.text assignment
        if not runs:
            para.text = desired_text
            return True

        # original run text lengths (to map slices)
        orig_run_texts = [r.text or "" for r in runs]
        orig_lengths = [len(t) for t in orig_run_texts]
        total_orig_len = sum(orig_lengths)

        # Overwrite runs slice-by-slice using original run lengths so run objects remain
        pos = 0
        for i, run in enumerate(runs):
            slot_len = orig_lengths[i]
            if pos >= len(desired_text):
                # no characters remain for this run
                if run.text != "":
                    run.text = ""
            else:
                end = min(pos + slot_len, len(desired_text))
                run_slice = desired_text[pos:end]
                run.text = run_slice
            pos += slot_len

        # If desired_text has extra characters beyond combined original run text, append to the last run
        if len(desired_text) > total_orig_len:
            extra = desired_text[total_orig_len:]
            runs[-1].text = (runs[-1].text or "") + extra

        return True

    def generate_noc(self, full_name: str, job_title: str, department: str, output_dir: str = "generated_noc") -> str:
        """
        Generates and saves a personalized NOC file by replacing placeholders in paragraphs and table cells.
        Returns the path of the generated file.
        """
        doc = Document(self.template_path)

        replacements = {
            "Full Name": full_name.strip(),
            "Job Title": job_title.strip(),
            "Department": department.strip()
        }

        # First pass: try run-preserving per-label replacements (safer for inline images/signatures)
        for para in doc.paragraphs:
            for key, value in replacements.items():
                if key.lower() in para.text.lower():
                    replaced = self._replace_field_in_paragraph_runs(para, key, value)
                    if not replaced:
                        # fallback: full-text replace (should be rare)
                        new_text = self._replace_fields_in_text(para.text, {key: value})
                        para.text = new_text

        # Process inside tables as well
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in replacements.items():
                            if key.lower() in para.text.lower():
                                replaced = self._replace_field_in_paragraph_runs(para, key, value)
                                if not replaced:
                                    new_text = self._replace_fields_in_text(para.text, {key: value})
                                    para.text = new_text

        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)
        safe_name = re.sub(r'[^A-Za-z0-9._-]+', '_', full_name.strip()) or "unknown"
        output_path = os.path.join(output_dir, f"NOC_{safe_name}.docx")

        doc.save(output_path)
        print(f"NOC generated successfully for {full_name}: {output_path}")
        return output_path


if __name__ == "__main__":
    generator = EmployeeNOCGenerator("NDA-1.docx")

    full_name = input("Enter Full Name: ").strip()
    job_title = input("Enter Job Title: ").strip()
    department = input("Enter Department: ").strip()

    generator.generate_noc(full_name, job_title, department)
